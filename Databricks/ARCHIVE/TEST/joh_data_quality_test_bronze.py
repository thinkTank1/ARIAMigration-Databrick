%pip install azure-storage-blob pandas faker python-docx

from pyspark.sql.functions import col, count, when, sum, isnan, isnull, SparkSession
from pyspark.sql.types import (
    StructType,
    StructField,
    StringType,
    IntegerType,
    DateType,
    BooleanType,
    FloatType,
)

from docx import Document

# Create a SparkSession
spark = SparkSession.builder \
    .appName("DataQualityTests") \
    .getOrCreate()

# Setting variables for use in subsequent cells
raw_mnt = "/mnt/ingest00rawsboxraw/ARIADM/ARM/JOH/test"
landing_mnt = "/mnt/ingest00landingsboxlanding/test"
bronze_mnt = "/mnt/ingest00curatedsboxbronze/ARIADM/ARM/JOH/test"
silver_mnt = "/mnt/ingest00curatedsboxsilver/ARIADM/ARM/JOH/test"
gold_mnt = "/mnt/ingest00curatedsboxgold/ARIADM/ARM/JOH/test"


def perform_data_quality_checks(df, table_name):
    """
    Performs data quality checks on the given DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the validation results.
    """
    validation_results = {}

    # Check if the table exists and has data
    validation_results[f"{table_name}_exists"] = df.count() > 0

    # Check if all required columns are present
    required_columns = df.columns
    missing_columns = [col for col in required_columns if col not in df.columns]
    validation_results[f"{table_name}_missing_columns"] = missing_columns

    # Check for null values in key columns
    key_columns = ["AdjudicatorId", "HistDate", "CentreId", "Role"]
    for column in key_columns:
        if column in df.columns:
            null_count = df.filter(col(column).isNull()).count()
            validation_results[f"{table_name}_{column}_null_count"] = null_count

    # Check for data type consistency
    schema = StructType(
        [
            StructField("AdjudicatorId", IntegerType(), nullable=False),
            StructField("Surname", StringType(), nullable=False),
            StructField("Forenames", StringType(), nullable=False),
            StructField("Title", StringType(), nullable=False),
            StructField("DateOfBirth", DateType(), nullable=True),
            StructField("CorrespondenceAddress", IntegerType(), nullable=True),
            StructField("ContactTelephone", StringType(), nullable=True),
            StructField("ContactDetails", StringType(), nullable=True),
            StructField("AvailableAtShortNotice", BooleanType(), nullable=True),
            StructField("DesignatedCentre", StringType(), nullable=True),
            StructField("EmploymentTerms", StringType(), nullable=True),
            StructField("FullTime", IntegerType(), nullable=True),
            StructField("IdentityNumber", StringType(), nullable=True),
            StructField("DateOfRetirement", DateType(), nullable=True),
            StructField("ContractEndDate", DateType(), nullable=True),
            StructField("ContractRenewalDate", DateType(), nullable=True),
            StructField("DoNotUse", BooleanType(), nullable=True),
            StructField("DoNotUseReason", StringType(), nullable=True),
            StructField("JudicialStatus", IntegerType(), nullable=True),
            StructField("Address1", StringType(), nullable=True),
            StructField("Address2", StringType(), nullable=True),
            StructField("Address3", StringType(), nullable=True),
            StructField("Address4", StringType(), nullable=True),
            StructField("Address5", StringType(), nullable=True),
            StructField("Postcode", StringType(), nullable=True),
            StructField("Telephone", StringType(), nullable=True),
            StructField("Mobile", StringType(), nullable=True),
            StructField("Email", StringType(), nullable=True),
            StructField("BusinessAddress1", StringType(), nullable=True),
            StructField("BusinessAddress2", StringType(), nullable=True),
            StructField("BusinessAddress3", StringType(), nullable=True),
            StructField("BusinessAddress4", StringType(), nullable=True),
            StructField("BusinessAddress5", StringType(), nullable=True),
            StructField("BusinessPostcode", StringType(), nullable=True),
            StructField("BusinessTelephone", StringType(), nullable=True),
            StructField("BusinessFax", StringType(), nullable=True),
            StructField("BusinessEmail", StringType(), nullable=True),
            StructField("JudicialInstructions", StringType(), nullable=True),
            StructField("JudicialInstructionsDate", DateType(), nullable=True),
            StructField("Notes", StringType(), nullable=True),
            StructField("HistDate", DateType(), nullable=False),
            StructField("HistType", IntegerType(), nullable=False),
            StructField("UserName", StringType(), nullable=True),
            StructField("Comment", StringType(), nullable=True),
            StructField("OtherCentres", StringType(), nullable=True),
            StructField("Role", IntegerType(), nullable=False),
            StructField("DateOfAppointment", DateType(), nullable=True),
            StructField("EndDateOfAppointment", DateType(), nullable=True),
        ]
    )
    data_type_mismatch_count = (
        df.select([count(when(~col(c).cast(t), 1)).alias(c) for c, t in schema.fields])
        .select(sum(col(c) for c in df.columns))
        .collect()[0][0]
    )
    validation_results[f"{table_name}_data_type_mismatch_count"] = (
        data_type_mismatch_count
    )

    return validation_results


# Perform data quality checks on the Bronze tables
bronze_tables = [
    ("bronze_adjudicator_et_hc_dnur", f"{bronze_mnt}/bronze_adjudicator_et_hc_dnur"),
    ("bronze_johistory_users", f"{bronze_mnt}/bronze_johistory_users"),
    ("bronze_othercentre_hearingcentre", f"{bronze_mnt}/bronze_othercentre_hearingcentre"),
    ("bronze_adjudicator_role", f"{bronze_mnt}/bronze_adjudicator_role"),
]

validation_results = {}

for table_name, table_path in bronze_tables:
    df = spark.read.format("delta").load(table_path)
    table_validation_results = perform_data_quality_checks(df, table_name)
    validation_results.update(table_validation_results)

# Create a new Word document
document = Document()

# Add a title to the document
document.add_heading('Bronze Data Quality Validation Report', level=1)

# Add a table to the document
table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
hdr_cells[2].text = 'Data Lineage'

# Add rows to the table
for metric, value in validation_results.items():
    table_name = metric.split("_")[0]
    column_name = "_".join(metric.split("_")[1:-1])
    lineage_info = f"{table_name}.{column_name}"
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(value)
    row_cells[2].text = lineage_info

# Save the document
document.save(f"{bronze_mnt}/reports/bronze_data_quality_validation_report.docx")
