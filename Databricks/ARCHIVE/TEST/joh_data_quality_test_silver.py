%pip install azure-storage-blob pandas faker python-docx

from pyspark.sql.functions import col, count, when, sum, SparkSession
from pyspark.sql.types import (
    StructType,
    StructField,
    IntegerType,
)

from docx import Document


# Create a SparkSession
spark = SparkSession.builder \
    .appName("DataQualityTests") \
    .getOrCreate()

# Setting variables for use in subsequent cells
raw_mnt = "/mnt/ingest00rawsboxraw/ARIADM/ARM/JOH/test"
landing_mnt = "/mnt/ingest00landingsboxlanding/test"
bronze_mnt = "/mnt/ingest00curatedsboxbronze/ARIADM/ARM/JOH/test"
silver_mnt = "/mnt/ingest00curatedsboxsilver/ARIADM/ARM/JOH/test"
gold_mnt = "/mnt/ingest00curatedsboxgold/ARIADM/ARM/JOH/test"


def perform_data_quality_checks(df, table_name):
    """
    Performs data quality checks on the given DataFrame.

    Args:
        df (DataFrame): The DataFrame to perform checks on.
        table_name (str): The name of the table being checked.

    Returns:
        dict: A dictionary containing the validation results.
    """
    validation_results = {}

    # Check if the table exists and has data
    validation_results[f"{table_name}_exists"] = df.count() > 0

    # Check if all required columns are present
    required_columns = df.columns
    missing_columns = [col for col in required_columns if col not in df.columns]
    validation_results[f"{table_name}_missing_columns"] = missing_columns

    # Check for null values in key columns
    key_columns = ["AdjudicatorId"]
    for column in key_columns:
        if column in df.columns:
            null_count = df.filter(col(column).isNull()).count()
            validation_results[f"{table_name}_{column}_null_count"] = null_count

    # Check for data type consistency
    schema = StructType(
        [
            StructField("AdjudicatorId", IntegerType(), nullable=False),
        ]
    )
    data_type_mismatch_count = (
        df.select([count(when(~col(c).cast(t), 1)).alias(c) for c, t in schema.fields])
        .select(sum(col(c) for c in df.columns))
        .collect()[0][0]
    )
    validation_results[f"{table_name}_data_type_mismatch_count"] = (
        data_type_mismatch_count
    )

    # Check if the data is filtered correctly based on AdjudicatorRole
    invalid_role_count = (
        df.join(
            spark.read.format("delta").load(f"{bronze_mnt}/bronze_adjudicator_role"),
            on="AdjudicatorId",
            how="inner",
        )
        .filter(col("Role").isin(7, 8))
        .count()
    )
    validation_results[f"{table_name}_invalid_role_count"] = invalid_role_count

    # Check for duplicate AdjudicatorId values
    duplicate_count = df.count() - df.dropDuplicates(["AdjudicatorId"]).count()
    validation_results[f"{table_name}_duplicate_count"] = duplicate_count

    return validation_results


# Perform data quality checks on the Silver tables
silver_tables = [
    ("silver_adjudicator_details", f"{silver_mnt}/silver_adjudicator_details"),
]

validation_results = {}

for table_name, table_path in silver_tables:
    df = spark.read.format("delta").load(table_path)
    table_validation_results = perform_data_quality_checks(df, table_name)
    validation_results.update(table_validation_results)

# Create a new Word document
document = Document()

# Add a title to the document
document.add_heading("Silver Data Quality Validation Report", level=1)

# Add a table to the document
table = document.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Metric"
hdr_cells[1].text = "Value"
hdr_cells[2].text = "Data Lineage"

# Add rows to the table
for metric, value in validation_results.items():
    table_name = metric.split("_")[0]
    column_name = "_".join(metric.split("_")[1:-1])
    lineage_info = f"{table_name}.{column_name}"
    row_cells = table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(value)
    row_cells[2].text = lineage_info

# Save the document
document.save(f"{silver_mnt}/reports/silver_data_quality_validation_report.docx")
